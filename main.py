from typing import Optional

import click as cli
from docx import Document
from docx.document import Document as DocumentObj
from docx.oxml.ns import qn
from docx.text.run import Run

from chars import replace_an_to_geo
from errors import BadFileExtension, UnsupportedWordDocFormat
from filters import AllTrueFilter, FontNameFilter, IFilter

WORD_FMT_SUPPORTED = {"docx"}
WORD_FMT_ALL = {"doc", "docx", "docm", "dot", "dotx", "dotm"}


def all_paragraphs_in_tables(doc: DocumentObj):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def iter_all_runs(doc: DocumentObj):
    # body
    for p in doc.paragraphs:
        for r in p.runs:
            yield r
    # tables
    for p in all_paragraphs_in_tables(doc):
        for r in p.runs:
            yield r
    # headers/footers
    for sec in doc.sections:
        for part in (sec.header, sec.footer):
            for p in part.paragraphs:
                for r in p.runs:
                    yield r
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                yield r


def resolve_font_name(run: Run) -> Optional[str]:
    if run.font and run.font.name:
        return run.font.name
    if run.style and run.style.font and run.style.font.name:
        return run.style.font.name
    rPr = run._r.rPr
    if rPr is not None and rPr.rFonts is not None:
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            v = rPr.rFonts.get(qn(f"w:{key}"))
            if v:
                return v
    return None


def replace_text_by_font(
    path_in: str,
    path_out: str,
    *,
    font_filter: Optional[IFilter] = None,
) -> int:
    if font_filter is None:
        font_filter = AllTrueFilter()

    doc = Document(path_in)

    changed_count = 0

    for run in iter_all_runs(doc):
        txt = run.text
        if not txt:
            continue

        f_name = resolve_font_name(run)
        if not font_filter.filter(f_name):
            continue

        new_text = replace_an_to_geo(txt)
        if new_text != txt:
            run.text = new_text
            changed_count += 1

    doc.save(path_out)
    return changed_count


def validate_filename(_0, _1, value: str) -> str:
    split_path = value.lower().rsplit(".", 1)
    ext = split_path[-1]
    if ext not in WORD_FMT_ALL:
        raise BadFileExtension(ext, value)
    if ext not in WORD_FMT_SUPPORTED:
        raise UnsupportedWordDocFormat(ext, value)
    return value


@cli.command()
@cli.option(
    "--input-file",
    "-i",
    type=str,
    required=True,
    help="Input file.",
    callback=validate_filename,
)
@cli.option(
    "--output-file",
    "-o",
    type=str,
    required=True,
    help="Output file.",
    callback=validate_filename,
)
def main(input_file: str, output_file: str):
    n = replace_text_by_font(
        input_file,
        output_file,
        font_filter=FontNameFilter("AcadNusx"),
    )
    print("#(Elements changed):", n)


if __name__ == "__main__":
    main()
